import os
import time
import random
import json
import re
import requests
import feedparser
from bs4 import BeautifulSoup
from newspaper import Article
import streamlit as st
from huggingface_hub import InferenceClient
import pandas as pd
from io import BytesIO
from docx import Document
import base64
import streamlit.components.v1 as components
from collections import defaultdict, Counter
from urllib.parse import urlparse, parse_qs
from datetime import date, timedelta  # 🆕 for date bucketing

# ======================
# Hugging Face Setup
# ======================
client = InferenceClient(
    provider="hf-inference",
    api_key=os.environ.get("xdpooja", "")
)

def summarize_text(text: str) -> str:
    if not text or not text.strip():
        return "No content to summarize."
    try:
        result = client.summarization(text, model="Falconsai/text_summarization")
        if hasattr(result, "summary_text"):
            return result.summary_text
        if isinstance(result, list) and result:
            first = result[0]
            if isinstance(first, dict) and "summary_text" in first:
                return first["summary_text"]
        if isinstance(result, dict) and "summary_text" in result:
            return result["summary_text"]
        return str(result)
    except Exception as e:
        return f"Error: {e}"

# ======================
# User-Agents
# ======================
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.3 Safari/605.1.15",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
]
def _UA():
    return {"User-Agent": random.choice(USER_AGENTS)}

# ======================
# DEFAULT CLUSTERS (fallback for Advanced mode)
# ======================
DEFAULT_CLUSTERS = {
    "Brand Monitoring": ["zyn","snus","tobacco pouch","nicotine pouch","marlboro","iqos"],
    "KOLs & Experts": ["kiran melkote","nimesh g desai","clive bates","rohan savio sequeira","david sweanor","bejon kumar misra","r zimlichman","upendra nath sharma","bharat gopal","chandrakant s pandav","mohsin wali"],
    "Regulatory & Institutional": ["icmr","indian council of medical research","who","world health organisation","fda","cdc","tobacco board","association of food scientists and technologists","confederation of indian food trade and industry","doctors against addiction","asian coalition of harm reduction","ache","cppr","the alternatives"],
    "Market Trends": ["cigarette","smoking","tobacco products","tobacco industry","tobacco use","big tobacco","traditional cigarettes","conventional cigarettes"],
    "Safer Alternatives (ENDS/HTP)": ["vape","vaping","electronic cigarette","e-cigarette","ends","heat-not-burn","heat not burn","heatnotburn","heated tobacco","electronically heated tobacco products","vaporizer","portable vaporizer","electric smoking system","smoke-free"],
    "NRT & Other Alternatives": ["nrt","nicotine replacement therapy","nicotine spray","snus","nicotine pouch","nicotine gum"],
    "Tobacco Control & Anti-Tobacco": ["tobacco control","anti-tobacco","tobacco burden","tobacco disease","healthcare expenditure"],
    "Economic & Policy": ["economy of tobacco","tobacco harm reduction","tobacco farmers","tobacco-growing","tobacco workers","tobacco control 3.0","human-centric approach to tobacco control"],
    "Country Comparisons": ["england vapes","england tobacco","new zealand tobacco","japan tobacco","sweden smoke-free"],
    "Philanthropy & Advocacy": ["bloomberg philanthropies","campaign for tobacco-free kids","vital strategies","pakistan bloomberg"],
}

# ======================
# FEEDS
# ======================
@st.cache_data
def fetch_feed(query: str, duration: int):
    """Quick mode: one RSS call (may cap ~100 items)."""
    rss_url = f"https://news.google.com/rss/search?q={requests.utils.quote(query)}%20when%3A{duration}d&hl=en-IN&gl=IN&ceid=IN:en"
    try:
        resp = requests.get(rss_url, timeout=12, headers=_UA())
        resp.raise_for_status()
        return feedparser.parse(resp.content).entries
    except Exception:
        return []

@st.cache_data
def fetch_feed_range(query: str, after_yyyy_mm_dd: str, before_yyyy_mm_dd: str):
    """One window using absolute dates (inclusive after, exclusive before)."""
    q = f'{query} after:{after_yyyy_mm_dd} before:{before_yyyy_mm_dd}'
    rss_url = f"https://news.google.com/rss/search?q={requests.utils.quote(q)}&hl=en-IN&gl=IN&ceid=IN:en"
    try:
        resp = requests.get(rss_url, timeout=12, headers=_UA())
        resp.raise_for_status()
        return feedparser.parse(resp.content).entries
    except Exception:
        return []

@st.cache_data
def fetch_feed_all(query: str, days: int, bucket_days: int = 7):
    """
    Exhaustive mode: slice the full period into date buckets and merge results.
    De-duplicates by link/id/title+source.
    """
    if days < 1:
        return []
    today = date.today()
    start_all = today - timedelta(days=days - 1)
    all_entries, seen = [], set()

    cur = start_all
    while cur <= today:
        win_end = min(cur + timedelta(days=bucket_days - 1), today)
        after_s = cur.isoformat()
        # 'before:' is exclusive -> add one day to include win_end
        before_s = (win_end + timedelta(days=1)).isoformat()

        batch = fetch_feed_range(query, after_s, before_s) or []
        for e in batch:
            # robust de-dup key
            link = getattr(e, "link", None) or (e.get("link") if isinstance(e, dict) else None)
            eid  = getattr(e, "id", None) or (e.get("id") if isinstance(e, dict) else None)
            title = getattr(e, "title", None) or (e.get("title") if isinstance(e, dict) else "")
            src = ""
            if isinstance(e, dict):
                s = e.get("source") or {}
                if isinstance(s, dict):
                    src = s.get("title", "")
            key = link or eid or f"{title}|{src}"
            if key not in seen:
                seen.add(key)
                all_entries.append(e)

        cur = win_end + timedelta(days=1)

    return all_entries

# ======================
# BASIC MODE: URL resolver & extractor
# ======================
def get_article_url_basic(rss_url: str) -> str:
    try:
        qs = parse_qs(urlparse(rss_url).query)
        if "url" in qs and qs["url"]:
            return qs["url"][0]
    except Exception:
        pass
    try:
        resp = requests.get(rss_url)
        resp.raise_for_status()
        data = BeautifulSoup(resp.text, 'html.parser').select_one('c-wiz[data-p]').get('data-p')
        obj = json.loads(data.replace('%.@.', '["garturlreq",'))
        payload = {'f.req': json.dumps([[['Fbv4je', json.dumps(obj[:-6] + obj[-2:]), 'null', 'generic']]])}
        headers = {'content-type': 'application/x-www-form-urlencoded;charset=UTF-8','user-agent': random.choice(USER_AGENTS)}
        response = requests.post("https://news.google.com/_/DotsSplashUi/data/batchexecute", headers=headers, data=payload)
        array_string = json.loads(response.text.replace(")]}'", ""))[0][2]
        return json.loads(array_string)[1]
    except Exception:
        return rss_url

def fetch_article_content_basic(url: str) -> str:
    try:
        article = Article(url)
        article.download()
        article.parse()
        return article.text
    except Exception:
        return "Content could not be extracted."

# ======================
# ADVANCED MODE: URL resolver (keeps your method + stronger fallbacks)
# ======================
def get_article_url_adv(rss_url: str) -> str:
    try:
        qs = parse_qs(urlparse(rss_url).query)
        if "url" in qs and qs["url"]:
            return qs["url"][0]
    except Exception:
        pass
    try:
        resp = requests.get(rss_url, headers=_UA(), timeout=12)
        resp.raise_for_status()
        data = BeautifulSoup(resp.text, 'html.parser').select_one('c-wiz[data-p]').get('data-p')
        obj = json.loads(data.replace('%.@.', '["garturlreq",'))
        payload = {'f.req': json.dumps([[['Fbv4je', json.dumps(obj[:-6] + obj[-2:]), 'null', 'generic']]])}
        headers = {'content-type': 'application/x-www-form-urlencoded;charset=UTF-8','user-agent': random.choice(USER_AGENTS)}
        response = requests.post("https://news.google.com/_/DotsSplashUi/data/batchexecute", headers=headers, data=payload, timeout=12)
        array_string = json.loads(response.text.replace(")]}'", ""))[0][2]
        return json.loads(array_string)[1]
    except Exception:
        pass
    try:
        r = requests.head(rss_url, allow_redirects=True, timeout=12, headers=_UA())
        if r.url and r.url.startswith("http") and "news.google.com" not in urlparse(r.url).netloc:
            return r.url
    except Exception:
        pass
    try:
        r = requests.get(rss_url, allow_redirects=True, timeout=12, headers=_UA())
        if r.url and r.url.startswith("http") and "news.google.com" not in urlparse(r.url).netloc:
            return r.url
    except Exception:
        pass
    return rss_url

def fetch_article_content_adv(url: str) -> str:
    try:
        a = Article(url); a.download(); a.parse()
        if a.text and a.text.strip():
            return a.text
    except Exception:
        pass
    try:
        r = requests.get(url, timeout=12, headers=_UA()); r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        art = soup.find("article")
        if art:
            text = "\n".join(p.get_text(" ", strip=True) for p in art.find_all("p"))
            if len(text) > 300:
                return text
        best = max(soup.find_all(["main","div","section"]), key=lambda tag: sum(len(p.get_text()) for p in tag.find_all("p")), default=None)
        if best:
            text = "\n".join(p.get_text(" ", strip=True) for p in best.find_all("p"))
            if len(text) > 150:
                return text
    except Exception:
        pass
    return "Content could not be extracted."

# ======================
# ADVANCED MODE: clusters + matching
# (unchanged)
# ======================
def split_terms(s: str):
    if not isinstance(s, str):
        s = str(s)
    s = s.replace("“", '"').replace("”", '"').replace("’", "'").strip().strip('"').strip("'")
    parts = re.split(r"\s*/\s*", s)
    out = []
    for p in parts:
        p = p.strip().strip('"').strip("'").replace("+", " ")
        p = re.sub(r"\s+", " ", p)
        if p:
            out.append(p)
    return out

def normalize_clusters(obj: dict) -> dict:
    clusters = defaultdict(set)
    if isinstance(obj, dict):
        for cat in obj.get("categories", []) or []:
            cat_name = (cat.get("name") or "").strip()
            subs = cat.get("subcategories")
            if isinstance(subs, list) and subs:
                for sub in subs:
                    sub_name = (sub.get("name") or "").strip() or cat_name or "General"
                    for kw in sub.get("keywords", []) or []:
                        for term in split_terms(kw):
                            clusters[sub_name].add(term.lower())
            else:
                for kw in cat.get("keywords", []) or []:
                    for term in split_terms(kw):
                        clusters[cat_name or "General"].add(term.lower())
        qcoh = (obj.get("queries") or {}).get("cohorts", {})
        if isinstance(qcoh, dict):
            for name, qstr in qcoh.items():
                q = str(qstr or "")
                phrases = re.findall(r'"([^"]+)"', q.replace("“","\"").replace("”","\""))
                tokens = []
                for token in ["ICMR","WHO","FDA","CDC","Tobacco Board","Snus","Zyn","IQOS","tobacco control","tobacco harm reduction","E-Cigarettes","Electronic cigarette","ENDS","Heat-Not-Burn","heatnotburn","Nicotine Pouch","NRT","Bloomberg Philanthropies"]:
                    if token.lower() in q.lower():
                        tokens.append(token)
                for item in phrases + tokens:
                    for term in split_terms(item):
                        clusters[f"Cohort: {name}"].add(term.lower())
        for k, v in obj.items():
            if isinstance(v, list) and all(not isinstance(x, dict) for x in v):
                for item in v:
                    for term in split_terms(item):
                        clusters[str(k)].add(term.lower())
    return {k: sorted(t for t in v if isinstance(t, str) and t.strip()) for k, v in clusters.items() if v}

def _flex(term: str) -> str:
    t = term.lower()
    t = t.replace("e-cigarette", r"e[ -]?cigarette")
    t = t.replace("heat-not-burn", r"heat[ -]?not[ -]?burn")
    t = t.replace("smoke-free", r"smoke[ -]?free")
    t = t.replace("nicotine pouch", r"nicotine[ -]?pouch")
    t = t.replace("heatnotburn", r"heat[ -]?not[ -]?burn")
    return t

def compile_patterns(cluster_dict: dict):
    compiled = {}
    for cluster, terms in (cluster_dict or {}).items():
        pats = []
        for t in terms or []:
            t = str(t or "").strip()
            if not t:
                continue
            pat = _flex(t)
            try:
                if re.match(r"^[a-z0-9\-\.]+$", t.lower()):
                    pats.append(re.compile(rf"\b{pat}\b", re.IGNORECASE))
                else:
                    pats.append(re.compile(pat, re.IGNORECASE))
            except re.error:
                continue
        if pats:
            compiled[cluster] = pats
    return compiled

def count_matches(text: str, pats) -> int:
    if not text:
        return 0
    total = 0
    for p in pats:
        try:
            total += len(p.findall(text))
        except Exception:
            total += 1 if p.search(text) else 0
    return total

def classify_article(title, source, link, article_text, compiled_clusters,
                     w_title=1.0, w_source=1.0, w_url=0.5, w_body=4.0):
    title = title or ""; source = source or ""; link = link or ""; body = article_text or ""
    scores = {}
    for cluster, pats in compiled_clusters.items():
        s = 0.0
        s += w_title  * count_matches(title,  pats)
        s += w_source * count_matches(source, pats)
        s += w_url    * count_matches(link,   pats)
        s += w_body   * count_matches(body,   pats)
        if s > 0:
            scores[cluster] = s
    if not scores:
        return None, 0.0, {}
    best = max(scores.values())
    tied = [k for k, v in scores.items() if v == best]
    primary = sorted(tied)[0]
    return primary, scores[primary], scores

# ======================
# EXPORTS (unchanged)
# ======================
def generate_excel(df: pd.DataFrame) -> bytes:
    buf = BytesIO()
    df_out = df.copy()
    if "MatchedClusters" in df_out.columns:
        df_out["MatchedClusters"] = df_out["MatchedClusters"].apply(
            lambda x: ", ".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
        )
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df_out.to_excel(writer, index=False, sheet_name="News")
    return buf.getvalue()

def generate_word_basic(df: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading("News Articles", 0)
    for _, row in df.iterrows():
        doc.add_heading(row.get('Headline', 'Untitled'), level=1)
        doc.add_paragraph(f"Source: {row.get('Source','')}")
        doc.add_paragraph(f"Date: {row.get('Published','')}")
        doc.add_paragraph(f"Link: {row.get('Link','')}")
        doc.add_paragraph("Article:", style="Intense Quote")
        doc.add_paragraph(str(row.get('Article', "")))
        doc.add_paragraph("Summary:", style="Intense Quote")
        doc.add_paragraph(str(row.get('Summary', "")))
        doc.add_paragraph("\n---\n")
    buf = BytesIO(); doc.save(buf); return buf.getvalue()

def generate_word_grouped(df: pd.DataFrame, top_n_per_cluster: int = 5) -> bytes:
    doc = Document()
    doc.add_heading("Clustered News Report", 0)
    if df.empty or "PrimaryCluster" not in df.columns:
        doc.add_paragraph("No clustered data available.")
        buf = BytesIO(); doc.save(buf); return buf.getvalue()
    for cluster in sorted(df["PrimaryCluster"].dropna().unique()):
        cluster_df = df[df["PrimaryCluster"] == cluster].copy()
        if cluster_df.empty:
            continue
        doc.add_heading(cluster, level=1)
        doc.add_heading(f"Top {top_n_per_cluster} Articles", level=2)
        top_articles = (cluster_df.sort_values(["RelevanceScore","Published"], ascending=[False, True]).head(top_n_per_cluster))
        for _, row in top_articles.iterrows():
            doc.add_heading(row.get("Headline","Untitled"), level=3)
            doc.add_paragraph(f"Source: {row.get('Source','')}")
            doc.add_paragraph(f"Date: {row.get('Published','')}")
            doc.add_paragraph(f"Link: {row.get('Link','')}")
            doc.add_paragraph("Summary:", style="Intense Quote")
            doc.add_paragraph(row.get("Summary","") or "-")
    buf = BytesIO(); doc.save(buf); return buf.getvalue()

# ======================
# UI — Header
# ======================
st.set_page_config(page_title="Mavs News Tracker", layout="wide")
colx, coly = st.columns(2)
with colx:
    st.image("Mavericks logo.png", width=150)
    st.title("Mavs News Tracker")

# ======================
# Sidebar — Mode toggle
# ======================
mode = st.sidebar.radio(
    "App mode",
    ["Basic v1", "Advanced v2 (clusters + scoring)"],
    index=1,
    help="Switch between basic script and the advanced classifier."
)

# Show clustering section only in Advanced mode
if mode == "Advanced v2 (clusters + scoring)":
    if "clusters" not in st.session_state:
        st.session_state.clusters = DEFAULT_CLUSTERS
    st.sidebar.subheader("🔧 Clustering Settings")
    uploaded = st.sidebar.file_uploader("Upload Cluster Config (JSON or CSV)", type=["json","csv"])
    if uploaded:
        try:
            if uploaded.type == "application/json" or uploaded.name.lower().endswith(".json"):
                cfg = json.load(uploaded)
                normalized = normalize_clusters(cfg)
                st.session_state.clusters = normalized if normalized else DEFAULT_CLUSTERS
                total_terms = sum(len(v) for v in st.session_state.clusters.values())
                st.sidebar.success(f"Loaded {len(st.session_state.clusters)} clusters / {total_terms} terms.")
                with st.sidebar.expander("Preview clusters", expanded=False):
                    st.json({k: v[:10] for k, v in list(st.session_state.clusters.items())[:10]})
            else:
                dfc = pd.read_csv(uploaded)
                cols = {c.lower().strip(): c for c in dfc.columns}
                cat_col = cols.get("category") or cols.get("cluster")
                term_col = cols.get("term")
                if not cat_col or not term_col:
                    st.sidebar.warning("CSV must include 'category' (or 'cluster') and 'term'. Using defaults.")
                else:
                    temp = defaultdict(set)
                    for _, r in dfc.iterrows():
                        cat = str(r.get(cat_col, "")).strip() or "General"
                        raw_term = r.get(term_col, "")
                        for t in split_terms(str(raw_term)):
                            if t.strip():
                                temp[cat].add(t.lower().strip())
                    parsed = {k: sorted(v) for k, v in temp.items() if v}
                    st.session_state.clusters = parsed if parsed else DEFAULT_CLUSTERS
                    total_terms = sum(len(v) for v in st.session_state.clusters.values())
                    st.sidebar.success(f"Loaded {len(st.session_state.clusters)} clusters / {total_terms} terms.")
                    with st.sidebar.expander("Preview clusters", expanded=False):
                        st.json({k: v[:10] for k, v in list(st.session_state.clusters.items())[:10]})
        except Exception as e:
            st.sidebar.error(f"Failed to parse config: {e}")

    st.sidebar.caption("This loader accepts UDAN-style JSON or simple {cluster:[terms]} JSON.")

# ======================
# Search Controls
# ======================
if "data" not in st.session_state:
    st.session_state.data = []

col1, col2, col3 = st.columns([2,1,1])
with col1:
    query = st.text_input("Search query", "Chemical Industry")
with col2:
    # allow up to 365 days
    duration = st.number_input("Duration (in days, up to 365)", min_value=1, max_value=365, value=1)
with col3:
    india_only = st.checkbox("India-focused filter", value=False, help="Keep articles that mention India/Indian")

# 🆕 Fetch-all controls
st.markdown("")
cfa1, cfa2 = st.columns([1,1])
with cfa1:
    fetch_all = st.checkbox("Fetch all (date-bucketed)", value=False, help="Pulls every RSS item in the period by slicing dates. May return far more than 100.")
with cfa2:
    bucket_days = st.number_input("Bucket size (days)", min_value=1, max_value=30, value=7, help="Smaller buckets can surface more items; 7 is a good balance.")

st.divider()

# ======================
# Run button (label adapts to mode)
# ======================
btn_label = "📡 Fetch News (Basic v1)" if mode.startswith("Basic") else "📡 Fetch & Classify News (Advanced v2)"
show_raw_article = st.checkbox("Show full article text in expanders", value=False)

if st.button(btn_label):
    # 🆕 choose fetch strategy
    if fetch_all:
        st.info("Fetching all articles across date buckets…")
        entries = fetch_feed_all(query, int(duration), int(bucket_days))
    else:
        entries = fetch_feed(query, int(duration))

    results = []
    progress = st.progress(0)
    status = st.empty()

    # advanced compilers if needed
    if mode.startswith("Advanced"):
        compiled = compile_patterns(st.session_state.get("clusters", DEFAULT_CLUSTERS))

    total = max(1, len(entries))
    for i, entry in enumerate(entries):
        progress.progress((i + 1) / total)
        status.text(f"Processing {i + 1}/{total}")

        title = getattr(entry, "title", None) or (entry.get("title") if isinstance(entry, dict) else "")
        source = ""
        if isinstance(entry, dict):
            src = entry.get("source") or {}
            if isinstance(src, dict):
                source = src.get("title", "")
        published = getattr(entry, "published", None) or (entry.get("published") if isinstance(entry, dict) else "") or getattr(entry, "updated", "")

        raw_link = getattr(entry, "link", None) or (entry.get("link") if isinstance(entry, dict) else "")

        if mode.startswith("Basic"):
            link = get_article_url_basic(raw_link) if raw_link else raw_link
            article = fetch_article_content_basic(link) if link else ""
        else:
            link = get_article_url_adv(raw_link) if raw_link else raw_link
            article = fetch_article_content_adv(link) if link else ""

        if india_only:
            haystack = " ".join([title or "", source or "", article or ""]).lower()
            if ("india" not in haystack) and ("indian" not in haystack):
                continue

        summary = summarize_text(article)

        row = {
            "Source": source,
            "Published": published,
            "Headline": title,
            "Link": link,
            "Article": article,
            "Summary": summary
        }

        if mode.startswith("Advanced"):
            primary_cluster, relevance, matches = classify_article(
                title, source, link, article, compiled,
                w_title=1.0, w_source=1.0, w_url=0.5, w_body=4.0
            )
            matched_names_sorted = [k for k, _ in sorted(matches.items(), key=lambda kv: kv[1], reverse=True)] if matches else []
            row.update({
                "PrimaryCluster": primary_cluster,
                "RelevanceScore": relevance,
                "MatchedClusters": matched_names_sorted
            })

        results.append(row)

        with st.expander(f"📰 {title}"):
            st.caption(f"Source: {source}")
            st.caption(f"Date: {published}")
            if link:
                st.markdown(f"[Read original article]({link})")
            if mode.startswith("Advanced"):
                st.caption(f"Primary Cluster: {row.get('PrimaryCluster','-')}  |  Score: {row.get('RelevanceScore',0):.1f}")
                mc = row.get("MatchedClusters") or []
                st.caption(f"Matched: {', '.join(mc) or '-'}")
            if show_raw_article:
                st.markdown("**Article:**"); st.write(article)
            st.markdown("**Summary:**"); st.write(summary)

        time.sleep(0.2 + random.random() * 0.4)

    st.session_state.data = results
    progress.empty()
    status.success("Done!")

# ======================
# Downloads
# ======================
if st.session_state.data:
    df = pd.DataFrame(st.session_state.data)
    st.subheader("💾 Download")

    excel_bytes = generate_excel(df)
    b64_xlsx = base64.b64encode(excel_bytes).decode("utf-8")
    c1, c2, c3 = st.columns(3)

    with c1:
        st.markdown(
            f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64_xlsx}" download="news.xlsx" style="text-decoration:none; padding:8px 16px; background-color:#4CAF50; color:white; border-radius:4px;">⬇️ Excel</a>',
            unsafe_allow_html=True,
        )

    if mode.startswith("Basic"):
        word_basic = generate_word_basic(df)
        b64_docx_basic = base64.b64encode(word_basic).decode("utf-8")
        with c2:
            st.markdown(
                f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx_basic}" download="news_basic.docx" style="text-decoration:none; padding:8px 16px; background-color:#1976D2; color:white; border-radius:4px;">⬇️ Word (basic)</a>',
                unsafe_allow_html=True,
            )
        with c3:
            st.caption("—")
    else:
        word_basic = generate_word_basic(df)
        b64_docx_flat = base64.b64encode(word_basic).decode("utf-8")
        with c2:
            st.markdown(
                f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx_flat}" download="news_flat.docx" style="text-decoration:none; padding:8px 16px; background-color:#1976D2; color:white; border-radius:4px;">⬇️ Word (flat)</a>',
                unsafe_allow_html=True,
            )
        by_cluster = df[df["PrimaryCluster"].notna()] if "PrimaryCluster" in df.columns else pd.DataFrame()
        if not by_cluster.empty:
            word_grouped = generate_word_grouped(by_cluster, top_n_per_cluster=5)
            b64_docx_grouped = base64.b64encode(word_grouped).decode("utf-8")
            with c3:
                st.markdown(
                    f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx_grouped}" download="news_by_cluster.docx" style="text-decoration:none; padding:8px 16px; background-color:#7B1FA2; color:white; border-radius:4px;">⬇️ Word (by cluster)</a>',
                    unsafe_allow_html=True,
                )
        else:
            with c3:
                st.caption("No clustered data yet.")

st.markdown("---")
st.caption("Toggle between Basic v1 and Advanced v2 from the sidebar.")
components.iframe("https://tally.so/embed/nPNe6Q?alignLeft=1&hideTitle=1&transparentBackground=1&dynamicHeight=1", width=400, height=500)
